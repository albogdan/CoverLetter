from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
import csv

TEMPLATE_NAME = 'CLTemplate.docx'


TOP_MARGIN = Inches(1)
BOTTOM_MARGIN = Inches(1)
LEFT_MARGIN = Inches(1)
RIGHT_MARGIN = Inches(1)

#LINE_SPACING = WD_LINE_SPACING.DOUBLE #Double spaced
LINE_SPACING = WD_LINE_SPACING.ONE_POINT_FIVE #Single spaced



#FONT_NAME = 'Calibri'
FONT_NAME = 'Times New Roman'
FONT_SIZE = Pt(11)


document = Document(TEMPLATE_NAME)


template = open("template.csv", "r")
keywords = template.readline().strip('\n').split(',')
print("Keywords: ", keywords)
csv_reader = csv.reader(template, delimiter=',')

for line in csv_reader:
    print(line)
    outDoc = Document()

    style = outDoc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = FONT_SIZE
    
    for para in document.paragraphs:
        text = para.text
        for i in range(0, len(keywords)):
            key = '{{' + keywords[i] + '}}'
            replacement = line[i]
            text = text.replace(key, replacement)
        paragraph = outDoc.add_paragraph(text)
        paragraph.style = outDoc.styles['Normal']
        paragraph.paragraph_format.line_spacing = LINE_SPACING
        
    for section in outDoc.sections:
        section.top_margin = TOP_MARGIN
        section.bottom_margin = BOTTOM_MARGIN
        section.left_margin = LEFT_MARGIN
        section.right_margin = RIGHT_MARGIN
    saveLocation = 'letters/' + str(line[keywords.index("Company".upper())]) + ' - ' + str(line[keywords.index("Position".upper())])+ ' CL.docx'
    outDoc.save(saveLocation)
    #print(para.text)
